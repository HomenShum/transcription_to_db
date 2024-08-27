import streamlit as st
from streamlit_webrtc import WebRtcMode, webrtc_streamer
import pydub
import time
import queue
import os
from twilio.rest import Client
import logging
from icecream import ic
from tenacity import retry, stop_after_attempt, wait_exponential
from openai import AzureOpenAI
from fastapi import HTTPException
from docx import Document
import requests
from st_login_form import login_form
from supabase import create_client
import base64
from cryptography.fernet import Fernet


logger = logging.getLogger(__name__)
st.session_state.audio_recorded = False

# Function to check if data is encrypted
def is_encrypted(data):
    try:
        base64.urlsafe_b64decode(data)
        return True
    except Exception:
        return False

# Function to encrypt data
def encrypt_data(data, key):
    fernet = Fernet(key)
    return fernet.encrypt(data.encode()).decode()

# Function to decrypt data
def decrypt_data(data, key):
    fernet = Fernet(key)
    return fernet.decrypt(data.encode()).decode()

# Initialize Supabase connection using st.cache_resource
@st.cache_resource
def init_connection():
    url = st.secrets["SUPABASE_URL"]
    key = st.secrets["SUPABASE_KEY"]
    logger.debug("Initializing Supabase connection")
    return create_client(url, key)

supabase = init_connection()

AzureOpenAI_client = AzureOpenAI(
    azure_endpoint=st.secrets["AOAIEndpoint"], 
    api_key=st.secrets["AOAIKey"],  
    api_version="2024-02-15-preview"
)

AzureOpenAIWhisper_client = AzureOpenAI(
    azure_endpoint=st.secrets["AOAIEndpointWhisper"],
    api_key=st.secrets["AOAIKey1Whisper"],
    api_version="2024-02-15-preview"
)

@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=2, max=6))
def request_chat_completion(user_prompt, system_prompt):
    response = AzureOpenAI_client.chat.completions.create(
        # model="gpt-4-0125-preview",
        model="gpt-4o",
        messages=[
            {"role":"system","content":system_prompt},
            {"role":"user","content":user_prompt}
        ]
    )
    if 'total_tokens' not in st.session_state:
        st.session_state.total_tokens = 0
    st.session_state.total_tokens += response.usage.total_tokens
    ic(st.session_state.total_tokens)

    return response.choices[0].message.content.replace('$', '\$')

def create_docx(transcriptions, summary):
    doc = Document()
    doc.add_heading('Summary', 0)
    doc.add_paragraph(summary)
    doc.add_heading('Transcriptions', 0)
    for transcription in transcriptions:
        doc.add_paragraph(transcription)
    doc.save('transcriptions_and_summary.docx')

@st.cache_data
def get_ice_servers():
    try:
        account_sid = st.secrets["TWILIO_ACCOUNT_SID"]
        auth_token = st.secrets["TWILIO_AUTH_TOKEN"]
    except KeyError:
        logger.warning("Twilio credentials are not set. Fallback to a free STUN server from Google.")
        return [{"urls": ["stun:stun.l.google.com:19302"]}]

    client = Client(account_sid, auth_token)
    token = client.tokens.create()
    return token.ice_servers

st.fragment(run_every=60)
def user_login():
    if "authenticated" not in st.session_state:
        st.session_state["authenticated"] = False

    if st.session_state["authenticated"] == False:
        client = login_form()

def main():
    # Check if user is logged in
    user_login()
    # Real-time transcription
    st.header("Real Time Speech-to-Text")
    app_sst()
    st.divider()


    # Transcription Summary and Download
    st.header("Transcription Summary and Download")
    transcribe_and_summarize()

def transcribe_audio(temp_wav_file, session_begin_timestamp_str, request_timestamp):
    logging.info(f"Transcribing audio file {temp_wav_file}")
    timestamp_str = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(request_timestamp))

    with open(temp_wav_file, "rb") as file:
        response = requests.post(
            st.secrets["transcription_api_url"],
            files={"file": file}
        )

    if response.status_code != 200:
        print(response.text)
        raise HTTPException(status_code=500, detail="Transcription failed")

    transcriptions = response.json()
    transcriptions_text = ' '.join([transcription for transcription in transcriptions["transcriptions"]])
    impending_result = f"{timestamp_str}: {transcriptions_text}"
    st.session_state.current_transcribed_text_and_timestamp_dict[session_begin_timestamp_str]['transcriptions'].append(impending_result.replace('$', '\$'))
    logging.info(f"Transcription for {temp_wav_file}: {impending_result}")


def process_audio_files(audio_files_queue, session_begin_timestamp_str):
    while not audio_files_queue.empty():
        temp_wav_file = audio_files_queue.get()
        transcribe_audio(temp_wav_file, session_begin_timestamp_str, time.time())
        os.remove(temp_wav_file)

@st.fragment
def app_sst():
    print("starting app_sst")
    
    webrtc_ctx = webrtc_streamer(
        key="speech-to-text",
        mode=WebRtcMode.SENDONLY,
        audio_receiver_size=40000,
        rtc_configuration={"iceServers": get_ice_servers()},
        media_stream_constraints={"video": False, "audio": True},
    )

    ic(webrtc_ctx)

    if not webrtc_ctx.state.playing:
        return

    if 'session_begin_timestamp_str' not in st.session_state:
        st.session_state.session_begin_timestamp_str = []
    session_begin_timestamp_str = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(time.time()))
    st.session_state.session_begin_timestamp_str.append(session_begin_timestamp_str)
    
    if 'current_transcribed_text_and_timestamp_dict' not in st.session_state:
        st.session_state.current_transcribed_text_and_timestamp_dict = {}

    st.session_state.current_transcribed_text_and_timestamp_dict[session_begin_timestamp_str] = {
        'transcriptions': [],
        'post_meeting_summary': "Press 'Summarize' to generate a summary."
    }

    audio_files_queue = queue.Queue()
    
    status_indicator = st.empty()

    with st.container(height=500, border=True):
        status_indicator.markdown("Loading...")
        display_transcriptions = st.empty()

    sound_chunk = pydub.AudioSegment.empty()
    last_ic_time = time.time()

    while True:
        if webrtc_ctx.audio_receiver:
            try:
                audio_frames = webrtc_ctx.audio_receiver.get_frames(timeout=1)
            except queue.Empty:
                time.sleep(0.1)
                status_indicator.markdown("No frame arrived.")
                continue

            status_indicator.markdown("Running. Say something!")

            for audio_frame in audio_frames:
                current_time = time.time()
                sound = pydub.AudioSegment(
                    data=audio_frame.to_ndarray().tobytes(),
                    sample_width=audio_frame.format.bytes,
                    frame_rate=audio_frame.sample_rate,
                    channels=len(audio_frame.layout.channels),
                )
                sound_chunk += sound

            if len(sound_chunk) > 0:
                if not st.session_state.audio_recorded:
                    st.session_state.audio_recorded = True
                    st.toast("Audio recorded successfully!")

            current_time = time.time()
            if current_time - last_ic_time >= 5: # Check if 5 seconds have passed

                # Ensure the directory exists
                os.makedirs("audio_chunks", exist_ok=True)


                # Export the sound_chunk to a temporary WAV file
                temp_wav_file = f"audio_chunks/temp_audio_{current_time}.wav"
                sound_chunk.export(temp_wav_file, format="wav")
                audio_files_queue.put(temp_wav_file)

                # Process the audio files synchronously
                process_audio_files(audio_files_queue, session_begin_timestamp_str)

                last_ic_time = current_time
                sound_chunk = pydub.AudioSegment.empty()

                try:
                    # Retrieve encryption key from Streamlit secrets
                    encryption_key = st.secrets["ENCRYPTION_KEY"]

                    # Encrypt the transcription text
                    encrypted_text = encrypt_data(st.session_state.current_transcribed_text_and_timestamp_dict[session_begin_timestamp_str]['transcriptions'][-1], encryption_key)

                    data = {
                        "username": st.session_state["username"],
                        "transcription_text": encrypted_text
                    }
                    response = supabase.table("transcriptions").insert(data).execute()
                    st.success("Transcription added successfully!")
                except Exception as e:
                    st.error(f"Failed to add transcription: {e}")


                jointed_transcription = "\n\n".join(st.session_state.current_transcribed_text_and_timestamp_dict[session_begin_timestamp_str]['transcriptions'])

                # Display the transcriptions
                display_transcriptions.markdown(f"**Transcriptions:**\n{jointed_transcription}")

            else:
                status_indicator.markdown("No audio to process.")
        else:
            status_indicator.markdown("Audio Receiver is not set. Abort.")
            break

@st.fragment
def transcribe_and_summarize():
    st.header("1. Format The Summary Template")
    ### TODO: Update the prompt value to the Patient Physician Conversation
    # value="""
    # Give me a half page note and summary report:
    # 0. Who are in the call and what's their roles?
    # 1. Company Overview (Concise 2-3 sentences); show revenue if available.
    # 2. What are their needs across the entire meeting?
    # 3. How much has the company raised?
    # 4. Who were the participants or investors to the fundraising?
    # 5. Future plans to fundraising?
    # 6. Current bank, reason, and any change of plans moving forward?
    # 7. Identify the NAICS code for this company
    # """, 
    default_system_prompt = st.text_area(
        label="What's your desired summary template?", 
        value = "Display the possible CPT and ICD codes. Summarize the conversation between the patient and the physician. Provide potential follow-up questions that can further enhance the diagnosis process.",
        height=300
    )

    if 'session_begin_timestamp_str' not in st.session_state:
        session_begin_timestamp_str = "Start a session to begin transcription."

    if 'session_begin_timestamp_str' in st.session_state:
        reversed_timestamp_list = st.session_state.session_begin_timestamp_str[::-1]
        session_begin_timestamp_str = st.selectbox(f"Select session begin timestamp:", reversed_timestamp_list)
        if session_begin_timestamp_str != "Start a session to begin transcription":
            if st.button("Summarize"):
                summary_request = request_chat_completion(system_prompt=default_system_prompt, user_prompt=str(st.session_state.current_transcribed_text_and_timestamp_dict[session_begin_timestamp_str]['transcriptions']))
                st.session_state.current_transcribed_text_and_timestamp_dict[session_begin_timestamp_str]["post_meeting_summary"] = summary_request            

    col1, col2 = st.columns(2)

    with col1:
        st.header("2. Transcriptions")
        if session_begin_timestamp_str != "Start a session to begin transcription.":
            st.subheader(f"Transcriptions for session begin timestamp: {session_begin_timestamp_str}")
            if 'current_transcribed_text_and_timestamp_dict' in st.session_state:
                for transcription in st.session_state.current_transcribed_text_and_timestamp_dict[session_begin_timestamp_str]['transcriptions']:
                    st.markdown(f"{transcription}")                                
            else:
                st.markdown("No transcriptions available.")

    with col2:
        st.header("3. Summary")
        if session_begin_timestamp_str != "Start a session to begin transcription.":
            st.subheader("Post Meeting Summary")
            if session_begin_timestamp_str in st.session_state.current_transcribed_text_and_timestamp_dict:
                if 'post_meeting_summary' in st.session_state.current_transcribed_text_and_timestamp_dict[session_begin_timestamp_str]:
                    if st.session_state.current_transcribed_text_and_timestamp_dict[session_begin_timestamp_str]["post_meeting_summary"] != "Press 'Summarize' to generate a summary.":
                        st.markdown(f"{st.session_state.current_transcribed_text_and_timestamp_dict[session_begin_timestamp_str]['post_meeting_summary']}")

    st.divider()
    if 'current_transcribed_text_and_timestamp_dict' in st.session_state:
        st.header("4. Download Transcriptions and Summary")
        create_docx(st.session_state.current_transcribed_text_and_timestamp_dict[session_begin_timestamp_str]['transcriptions'], st.session_state.current_transcribed_text_and_timestamp_dict[session_begin_timestamp_str]['post_meeting_summary'])
        st.download_button(
            label="Download Transcriptions and Summary",
            data=open("transcriptions_and_summary.docx", "rb").read(),
            file_name="transcriptions_and_summary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()
